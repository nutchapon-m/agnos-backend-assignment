#!/usr/bin/env python3
"""Compile the docs/*.md sources into one Word document.

Usage:
    python3 -m pip install python-docx    # once
    python3 docs/build-docx.py

The markdown files are the source of truth; agnos-backend-assignment-docs.docx is
generated. A hand edit in Word is lost on the next run.

Supported markdown: ATX headings, paragraphs, bullet/numbered lists (with indented
continuation lines), GFM pipe tables, fenced code blocks, blockquotes, thematic
breaks, and inline code / bold / italic / strikethrough / links.
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parent
OUTPUT = DOCS_DIR / "agnos-backend-assignment-docs.docx"

# Sections, in reading order: (markdown file, printed title).
SECTIONS = [
    ("development-plan.md", "Development Plan"),
    ("er-diagram.md", "Data Model (ER Diagram)"),
    ("api-spec.md", "API Specification"),
]

TITLE = "Hospital Middleware Backend"
SUBTITLE = "Technical Documentation"
META = [
    "Repository: agnos-backend-assignment",
    "Stack: Go, Gin, sqlx, PostgreSQL, golang-migrate, nginx",
]

# --- styling ----------------------------------------------------------------

ACCENT = RGBColor(0x1A, 0x5F, 0x8A)  # headings, title
MUTED = RGBColor(0x5F, 0x63, 0x68)  # captions, links, code labels
CODE_INK = RGBColor(0xB0, 0x30, 0x60)  # inline code
CODE_FILL = "F2F3F5"  # code block shading
MONO = "Menlo"

BODY_SIZE = Pt(10.5)
CODE_INLINE_SIZE = Pt(9.5)
TABLE_SIZE = Pt(9.5)
TABLE_CODE_SIZE = Pt(8.5)
CODE_BLOCK_SIZE = Pt(8.5)
CAPTION_SIZE = Pt(8.5)

BODY_AFTER = Pt(6)
LIST_AFTER = Pt(2)


# --- markdown model ---------------------------------------------------------


@dataclass
class Block:
    kind: str  # heading | para | bullet | number | code | table | quote | rule
    text: str = ""
    level: int = 0
    lang: str = ""
    lines: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


FENCE_RE = re.compile(r"^```\s*([A-Za-z0-9+_-]*)\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*$")
BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
NUMBER_RE = re.compile(r"^(\d+)\.\s+(.*)$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
RULE_RE = re.compile(r"^-{3,}\s*$")
TABLE_SEP_RE = re.compile(r"^\|[\s:|-]+\|$")


def split_row(line: str) -> list[str]:
    """Split a pipe-table row, ignoring pipes inside code spans."""
    cells, buf, in_code = [], [], False
    for ch in line.strip().strip("|"):
        if ch == "`":
            in_code = not in_code
            buf.append(ch)
        elif ch == "|" and not in_code:
            cells.append("".join(buf).strip())
            buf = []
        else:
            buf.append(ch)
    cells.append("".join(buf).strip())
    return cells


def parse(md: str) -> list[Block]:
    lines = md.splitlines()
    blocks: list[Block] = []
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        fence = FENCE_RE.match(line)
        if fence:
            lang, body = fence.group(1), []
            i += 1
            while i < n and not FENCE_RE.match(lines[i]):
                body.append(lines[i])
                i += 1
            i += 1  # closing fence
            blocks.append(Block("code", lang=lang, lines=body))
            continue

        heading = HEADING_RE.match(line)
        if heading:
            blocks.append(
                Block("heading", text=heading.group(2), level=len(heading.group(1)))
            )
            i += 1
            continue

        if RULE_RE.match(line):
            blocks.append(Block("rule"))
            i += 1
            continue

        # table: a header row followed by a separator row
        if line.lstrip().startswith("|") and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1].strip()):
            rows = [split_row(line)]
            i += 2
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            blocks.append(Block("table", rows=rows))
            continue

        quote = QUOTE_RE.match(line)
        if quote:
            parts = [quote.group(1)]
            i += 1
            while i < n and (m := QUOTE_RE.match(lines[i])):
                parts.append(m.group(1))
                i += 1
            blocks.append(Block("quote", text=" ".join(p.strip() for p in parts).strip()))
            continue

        item = BULLET_RE.match(line) or NUMBER_RE.match(line)
        if item:
            kind = "bullet" if BULLET_RE.match(line) else "number"
            text = item.group(1) if kind == "bullet" else item.group(2)
            parts = [text]
            i += 1
            # continuation lines are indented and not a new item
            while i < n and lines[i].strip() and lines[i][:1] in " \t":
                stripped = lines[i].strip()
                if BULLET_RE.match(stripped) or NUMBER_RE.match(stripped):
                    break
                parts.append(stripped)
                i += 1
            blocks.append(Block(kind, text=" ".join(parts)))
            continue

        # paragraph: consecutive plain lines
        parts = [line.strip()]
        i += 1
        while i < n and lines[i].strip():
            nxt = lines[i]
            if (
                FENCE_RE.match(nxt)
                or HEADING_RE.match(nxt)
                or RULE_RE.match(nxt)
                or QUOTE_RE.match(nxt)
                or nxt.lstrip().startswith("|")
                or BULLET_RE.match(nxt)
                or NUMBER_RE.match(nxt)
            ):
                break
            parts.append(nxt.strip())
            i += 1
        blocks.append(Block("para", text=" ".join(parts)))

    return blocks


# --- inline runs ------------------------------------------------------------

INLINE_RE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<link>\[[^\]]+\]\([^)]*\))"
    r"|(?P<bold>\*\*.+?\*\*)"
    r"|(?P<strike>~~.+?~~)"
    r"|(?P<italic>(?<![\w*])\*(?![\s*])[^*]+\*)"
)

# A run is (text, flags), where flags is a set drawn from
# {"code", "bold", "italic", "strike", "link"}.
Run = tuple


def inline_runs(text: str, flags: frozenset = frozenset()) -> list[Run]:
    """Split text into styled runs, honouring nesting (bold around code, …)."""
    out: list[Run] = []
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            out.append((text[pos : m.start()], flags))
        kind, raw = m.lastgroup, m.group()
        if kind == "code":
            out.append((raw[1:-1], flags | {"code"}))
        elif kind == "link":
            label = raw[1 : raw.index("](")]
            out.extend(inline_runs(label, flags | {"link"}))
        elif kind == "bold":
            out.extend(inline_runs(raw[2:-2], flags | {"bold"}))
        elif kind == "strike":
            out.extend(inline_runs(raw[2:-2], flags | {"strike"}))
        else:
            out.extend(inline_runs(raw[1:-1], flags | {"italic"}))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], flags))
    return [(t, f) for t, f in out if t]


def add_runs(paragraph, text, *, body_size=BODY_SIZE, code_size=CODE_INLINE_SIZE):
    for chunk, flags in inline_runs(text):
        run = paragraph.add_run(chunk)
        if "code" in flags:
            run.font.name = MONO
            run.font.size = code_size
            run.font.color.rgb = CODE_INK
        else:
            run.font.size = body_size
        run.bold = "bold" in flags or None
        run.italic = ("italic" in flags or "link" in flags) or None
        if "strike" in flags:
            run.font.strike = True
        if "link" in flags and "code" not in flags:
            run.font.color.rgb = MUTED
        if "strike" in flags and "code" not in flags:
            run.font.color.rgb = MUTED
    if not paragraph.runs:
        paragraph.add_run("")


# --- docx emission ----------------------------------------------------------


def shade(cell, fill: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def page_break(doc) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def emit_code(doc, block: Block) -> None:
    if block.lang:
        label = doc.add_paragraph()
        label.paragraph_format.space_after = Pt(0)
        run = label.add_run(block.lang.upper())
        run.font.name = MONO
        run.font.size = Pt(7.5)
        run.font.color.rgb = MUTED
        run.bold = True

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade(cell, CODE_FILL)

    lines = block.lines or [""]
    for idx, line in enumerate(lines):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.font.name = MONO
        run.font.size = CODE_BLOCK_SIZE
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def emit_table(doc, block: Block) -> None:
    header, *body = block.rows
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    for cell, text in zip(table.rows[0].cells, header):
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        for chunk, flags in inline_runs(text):
            run = para.add_run(chunk)
            run.bold = True
            if "code" in flags:
                run.font.name = MONO
                run.font.size = TABLE_CODE_SIZE
            else:
                run.font.size = TABLE_SIZE

    for row in body:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_runs(para, text, body_size=TABLE_SIZE, code_size=TABLE_CODE_SIZE)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def emit_block(doc, block: Block) -> None:
    if block.kind == "heading":
        # markdown h1 is the document title, already emitted; the rest map 1:1
        level = min(block.level, 4)
        para = doc.add_heading(level=level)
        for chunk, flags in inline_runs(block.text):
            run = para.add_run(chunk)
            run.font.color.rgb = ACCENT
            if "code" in flags:
                run.font.name = MONO
        return

    if block.kind == "para":
        para = doc.add_paragraph()
        para.paragraph_format.space_after = BODY_AFTER
        add_runs(para, block.text)
        return

    if block.kind in ("bullet", "number"):
        style = "List Bullet" if block.kind == "bullet" else "List Number"
        para = doc.add_paragraph(style=style)
        para.paragraph_format.space_after = LIST_AFTER
        add_runs(para, block.text)
        return

    if block.kind == "quote":
        para = doc.add_paragraph()
        para.paragraph_format.space_after = BODY_AFTER
        para.paragraph_format.left_indent = Inches(0.3)
        add_runs(para, block.text)
        for run in para.runs:
            if run.font.name != MONO:
                run.italic = True
                run.font.color.rgb = MUTED
        return

    if block.kind == "code":
        emit_code(doc, block)
        return

    if block.kind == "table":
        emit_table(doc, block)
        return

    if block.kind == "rule":
        doc.add_paragraph().paragraph_format.space_after = Pt(8)


def title_page(doc) -> None:
    for _ in range(3):
        doc.add_paragraph()

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(SUBTITLE)
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED

    doc.add_paragraph()

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, line in enumerate(META):
        if idx:
            para.add_run("\n")
        run = para.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED


def contents(doc, parsed: list[tuple[str, str, list[Block]]]) -> None:
    heading = doc.add_heading(level=1)
    heading.add_run("Contents").font.color.rgb = ACCENT

    for number, (filename, title, blocks) in enumerate(parsed, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(f"{number}.  {title}")
        run.bold = True
        run.font.size = Pt(11.5)
        run = para.add_run(f"    ({filename})")
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

        for block in blocks:
            if block.kind == "heading" and block.level == 2:
                entry = doc.add_paragraph()
                entry.paragraph_format.space_after = Pt(0)
                entry.paragraph_format.left_indent = Inches(0.4)
                clean = re.sub(r"[`*]", "", block.text)
                entry.add_run(clean).font.size = Pt(10)


def build() -> Path:
    parsed: list[tuple[str, str, list[Block]]] = []
    for filename, title in SECTIONS:
        path = DOCS_DIR / filename
        if not path.exists():
            sys.exit(f"missing source: {path}")
        parsed.append((filename, title, parse(path.read_text(encoding="utf-8"))))

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(0.79))  # 2 cm

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = BODY_SIZE

    title_page(doc)
    page_break(doc)
    contents(doc, parsed)

    for number, (filename, title, blocks) in enumerate(parsed, start=1):
        page_break(doc)

        heading = doc.add_heading(level=1)
        heading.add_run(f"{number}. {title}").font.color.rgb = ACCENT

        source = doc.add_paragraph()
        source.paragraph_format.space_after = Pt(10)
        run = source.add_run(f"Source: docs/{filename}")
        run.italic = True
        run.font.size = CAPTION_SIZE
        run.font.color.rgb = MUTED

        for block in blocks:
            if block.kind == "heading" and block.level == 1:
                continue  # the file's own title, replaced above
            emit_block(doc, block)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build()
    print(f"wrote {out.relative_to(Path.cwd()) if out.is_relative_to(Path.cwd()) else out}")
